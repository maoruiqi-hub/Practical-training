from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = Path("specs/黄榆航/黄榆航-软件系统开发实践报告.docx")


def set_run_font(run, size=12, bold=False, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_font(paragraph, size=12, bold=False):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)


def add_paragraph(doc, text="", size=12, bold=False, align=None, first_line=True):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    if first_line and text:
        pf.first_line_indent = Pt(24)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.first_line_indent = Pt(-12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("• " + text)
    set_run_font(run, size=12)
    return p


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=180, bottom=80, end=180):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn("w:" + name))
        if element is None:
            element = OxmlElement("w:" + name)
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_cover(doc):
    add_paragraph(doc, "课程编号：A070848", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(12)
    r = title.add_run("软件系统开发实践\n实践报告")
    set_run_font(r, size=18, bold=True)

    doc.add_paragraph()
    rows = [
        ("姓名", "黄榆航"),
        ("学号", "20246153"),
        ("班级", "软件2409"),
        ("指导教师", "姜琳颖"),
        ("实践课程名称", "软件系统开发实训"),
        ("开设学期", "2024-2025 春季学期"),
        ("开设时间", ""),
        ("报告日期", ""),
    ]
    for left, right in rows:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.5
        r1 = p.add_run(f"{left}：")
        set_run_font(r1, size=12, bold=True)
        r2 = p.add_run(right)
        set_run_font(r2, size=12, bold=False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(96)
    r = p.add_run("东北大学软件学院")
    set_run_font(r, size=14, bold=True)
    doc.add_page_break()


def add_body(doc):
    add_heading1(doc, "实践目的 Aim of Project")
    add_paragraph(
        doc,
        "本次软件系统开发实践的目标是围绕“AI智慧课程平台”完成一个可运行、可扩展的教学辅助系统。项目希望在课程管理、任务提交、测评评价、能力图谱、学情分析等基础功能之上，引入 AI 辅助评价和学习诊断能力，为教师教学管理和学生学习反馈提供支持。",
    )
    add_paragraph(
        doc,
        "我在本次实践中主要承担项目基础版本打底和部分核心模块建设工作，包括数据库基础版本设计、软件前后端架构初始搭建、国产数据库适配与数据迁移、教师端前端界面设计与功能整合维护，以及系统模块三“测评与成果评价模块”的设计和开发。通过这些工作，保障项目具备后续扩展、集成和迭代的基础。",
    )

    add_heading1(doc, "预习内容 Preview Content")
    add_paragraph(
        doc,
        "在实践开始前，我预习和梳理了 Web 应用开发、前后端分离架构、关系型数据库建模、RESTful API 设计、教学平台业务流程和国产数据库基本使用方式等内容。前端方面主要关注页面路由、组件组织、接口封装和教师端交互流程；后端方面主要关注控制层、服务层、数据访问层的分层结构，以及实体、接口和数据库表之间的对应关系。",
    )
    add_paragraph(
        doc,
        "数据库方面，我重点预习了用户、课程、班级、任务、提交、知识点等教学业务对象之间的关系，理解一门课程从创建、分配、发布任务、学生提交到教师评价的基本数据流。国产数据库应用方面，我了解了项目从通用关系型数据库向 KingbaseES 等国产数据库适配时需要关注的连接配置、数据类型兼容、SQL 语法差异和数据迁移验证问题。",
    )

    add_heading1(doc, "实践内容和实践过程 Project Content and Process")

    add_heading2(doc, "3.1 概述 Overview")
    add_paragraph(
        doc,
        "本项目为 AI智慧课程平台，面向教师和学生两类用户，提供课程管理、任务管理、在线测验、题库管理、学习分析、能力图谱和 AI 辅助评价等功能。项目整体采用前后端分离方式开发，前端承担教师端、学生端页面展示和交互，后端提供统一业务接口和数据处理能力，数据库保存课程教学过程中的核心数据。",
    )
    add_paragraph(
        doc,
        "我在项目早期主要负责搭建基础版本，使系统先具备可扩展的整体框架。具体包括：设计用户、课程、班级、任务、提交、知识点等基础数据关系；搭建前后端基础结构和接口调用方式；完成国产数据库的应用、数据搬运和验证；持续维护教师端页面，将后续分散实现的功能整合到教师工作流中；并围绕模块三“测评与成果评价模块”推进题库、组卷、在线测验、自动评阅和评价数据沉淀等工作。需要说明的是，当前项目最终版本已经经过团队多轮迭代和改进，我负责的是其中的基础设计、模块实现和持续整合部分。",
    )

    add_heading2(doc, "3.2 相关技术 Relevant Technologies")
    add_paragraph(
        doc,
        "本项目使用前后端分离架构。前端以 Vue 技术栈为基础，按照页面、组件和接口封装进行组织，教师端页面包含课程列表、成绩统计、学情分析、班级运营、能力图谱、学生画像、题库与任务管理等功能入口。后端采用 Java Web 技术栈提供 RESTful API，通过 Controller、Service、Mapper 等层次组织业务逻辑和数据访问。前后端通过 JSON 数据交互，保证页面功能与后端业务接口相互解耦。",
    )
    add_paragraph(
        doc,
        "数据库设计采用关系型数据库模型，围绕课程教学场景抽象用户、教师、学生、课程、班级、任务、题目、提交、作答明细、知识点等实体。国产数据库应用方面，项目使用 KingbaseES 进行适配和验证，重点处理数据库连接配置、已有数据迁移、表结构兼容和查询结果验证等工作。AI 相关能力则通过独立服务或接口进行接入，避免 AI 能力和基础业务流程过度耦合。",
    )

    add_heading2(doc, "3.3 系统设计 System Design")
    add_paragraph(
        doc,
        "在系统设计阶段，我首先完成了基础数据模型的梳理。用户、课程、班级、任务、提交、知识点都属于我参与打底的核心对象。其中，用户和角色用于区分教师、学生和管理员；课程与班级用于表达教学组织关系；任务用于承载作业、实验报告和在线测验；提交用于记录学生完成情况；知识点用于支撑题库、测评分析、能力图谱和学习推荐等后续功能。",
    )
    add_paragraph(
        doc,
        "在前后端架构方面，我搭建了最基础的软件版本，使后续成员可以在统一框架上继续扩展。基础版本的重点不是一次性完成所有功能，而是明确前端页面组织、后端接口分层、数据库实体关系和接口调用规范，为后续模块开发提供稳定入口。这样可以避免不同功能各自独立堆叠，保证系统后续仍能整合为一个完整平台。",
    )
    add_paragraph(
        doc,
        "在模块三“测评与成果评价模块”设计中，我依据需求将模块范围划分为题库管理、试卷生成、在线测验、客观题评阅、主观成果评价和错题分析。该模块的核心目标是回答“如何评价学生学习成果”，其产生的测评数据和评价结果还要继续服务学生画像、学习推荐和教师学情分析。因此，设计时不仅关注题目和考试本身，也关注提交后的作答明细、知识点得分、错题统计和教师复核流程。",
    )
    add_paragraph(
        doc,
        "在国产数据库应用方面，我参与了数据库的迁移和适配工作，将项目数据从原有环境搬运到国产数据库环境中，并通过实际查询和系统运行验证数据可用性。该部分工作保证系统不只是理论上支持国产数据库，而是在项目展示和运行过程中确实使用了国产数据库环境。",
    )

    add_heading2(doc, "3.4 编码 Coding")
    add_paragraph(
        doc,
        "编码过程中，我首先完成了系统基础版本的前后端打底。前端方面，主要搭建教师端页面的基础结构、导航入口和功能区域，为课程、任务、成绩、学情、能力图谱、学生画像等页面提供统一入口。后端方面，围绕基础业务对象建立实体、接口和服务逻辑，使前端能够通过接口获取课程、任务、学生提交和评价相关数据。",
    )
    add_paragraph(
        doc,
        "在教师端开发中，我主要负责教师端前端界面设计、功能整合和维护。教师端在系统中偏向管理、诊断和干预角色，因此页面设计上需要把分散功能组织成教师可以理解的工作流。例如，教师需要能够进入课程、查看任务和提交、查看班级学习情况、进行成绩统计、进入题库和测评相关功能。后续项目增加 AI 能力、学情分析和能力图谱后，我也参与了教师端入口整合，使这些能力能够出现在合适的教师操作场景中。",
    )
    add_paragraph(
        doc,
        "在模块三编码中，我围绕题库管理、基础组卷、在线测验、客观题评阅和测评数据保存推进实现。题库部分支持多题型和课程维度管理；组卷部分支持基础随机组卷、按知识点组卷和难度相关配置；在线测验部分支持学生答题与提交；客观题评阅部分实现单选、多选等题型的基础自动评分；提交和作答数据则为后续错题分析、知识点掌握统计和教师复核提供数据基础。需要说明的是，最终版本中部分能力已经由团队继续改进，我的工作重点是完成基础链路和关键模块设计开发。",
    )

    add_heading2(doc, "3.5 测试 Testing")
    add_paragraph(
        doc,
        "测试过程中，我主要进行的是系统层面的手动测试，尤其是围绕教师端模拟真实教学场景和教师操作行为，而不是单元测试或专门的性能测试平台测试。我会以教师视角反复操作系统，例如进入课程、查看班级、发布任务、进入题库、生成测验、查看学生提交、检查成绩统计和学情分析入口，观察这些流程是否能够顺利连贯地完成。",
    )
    add_paragraph(
        doc,
        "在具体测试中，我重点检查数据库表关系是否能够支撑真实业务流程，例如课程下是否能发布任务，学生提交数据是否能被教师端读取，题目、任务和知识点之间是否能建立基本关联。对于教师端功能，我主要检查页面入口是否完整、接口数据是否能正确渲染、异常数据是否会影响页面打开，以及多个功能连续点击、切换页面时是否会出现明显问题。对于模块三，我主要通过手动流程验证题库、组卷、测验提交和客观题评阅是否形成闭环。测试中也暴露出一些问题，例如 AI 功能响应时间较长、部分边界情况处理不足、部分功能在连续操作时稳定性不够，这些都为后续优化提供了方向。",
    )

    add_heading1(doc, "实践总结 Conclusion")
    add_paragraph(
        doc,
        "通过本次实践，我对一个教学平台从基础数据模型、前后端架构、数据库适配到具体业务模块开发的完整过程有了更清晰的认识。项目早期的基础版本虽然功能较简单，但它决定了后续系统能否顺利扩展。用户、课程、班级、任务、提交、知识点等基础关系如果设计不清楚，后续 AI 评价、能力图谱、学情分析等功能都会缺少稳定的数据基础。",
    )
    add_paragraph(
        doc,
        "在教师端开发和维护过程中，我也认识到教师端不能只堆叠功能，而应该围绕教师真实工作流组织页面。教师更需要快速看到课程运行状态、学生提交情况、班级风险和需要干预的问题。因此，教师端的价值在于管理、诊断、校准和干预，而学生端则更需要后续重点打磨学习反馈和成长引导体验。",
    )
    add_paragraph(
        doc,
        "在国产数据库应用过程中，我体会到数据库适配不仅是修改连接配置，还包括表结构兼容、数据迁移、查询验证和系统运行验证。只有系统实际连接并使用国产数据库，才能证明适配工作真正完成。模块三的开发也让我认识到，测评模块不仅是出题和判分，更重要的是把评价数据沉淀下来，继续服务学习分析、错题诊断和能力画像。",
    )
    add_paragraph(
        doc,
        "本次实践也暴露出项目中的不足。由于系统在短时间内快速迭代，部分代码存在边界处理不充分、功能耦合较高、AI 响应较慢等问题。后续如果继续完善，应进一步提升系统稳定性，增强 AI 评价可信度，完善能力图谱与岗位能力之间的映射，并将系统定位为轻量、可插拔的 AI 教学增强工具，而不是臃肿的大型教学平台。",
    )

    add_heading1(doc, "参考资料 References")
    refs = [
        "项目组，《AI智慧课程平台项目需求文档》。",
        "项目组，《测评与成果评价模块差距分析与实施计划》。",
        "KingbaseES 官方文档：数据库连接、SQL 兼容与应用迁移相关说明。",
        "Vue、Spring Boot、RESTful API 与关系型数据库设计相关课程资料。"
    ]
    for ref in refs:
        add_paragraph(doc, ref, first_line=False)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    add_cover(doc)
    add_body(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
